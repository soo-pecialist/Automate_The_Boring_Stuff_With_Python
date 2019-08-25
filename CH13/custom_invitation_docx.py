#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Sat Aug 24 23:57:11 2019

@author: Soo Hyeon Kim
Read guest names from text file and generate a Word document with custom 
invitations.
"""

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

with open('guest.txt') as f:
    names = f.readlines()
    document = docx.Document()
    
    for name in names:
        name = name.strip()
        
        paragraph = document.add_paragraph('It would be a pleasure to have the' 
                               ' company of', style='Heading1')
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        paragraph = document.add_paragraph(name, style='Caption')
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        paragraph = document.add_paragraph('at 11010 Memory Lane on '
                                           'the Evening of', style='Heading1')
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        paragraph = document.add_paragraph('April 1st', style='BodyText')
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        paragraph = document.add_paragraph('at 7 o\'clock', style='Heading1')
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        document.add_page_break()
    
    document.save('invites.docx')
    
    print("'invites.docx' has been created")
